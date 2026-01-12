"""
RAG (Retrieval-Augmented Generation) 服务
为项目提供文档索引、检索和增强生成功能
"""

import os
import logging
import hashlib
import asyncio
from typing import List, Dict, Any, Optional, AsyncGenerator, Set
from pathlib import Path
from datetime import datetime
import json

# Word 文档支持
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed, Word document support disabled")

logger = logging.getLogger("RAGService")

try:
    import chromadb
    from chromadb.config import Settings
    from chromadb.utils import embedding_functions
    CHROMADB_AVAILABLE = True
except ImportError:
    logger.warning("chromadb not available. Install with: pip install chromadb")
    CHROMADB_AVAILABLE = False

try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    logger.warning("sentence-transformers not available. Install with: pip install sentence-transformers")
    SENTENCE_TRANSFORMERS_AVAILABLE = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
    SKLEARN_AVAILABLE = True
except ImportError:
    logger.warning("scikit-learn not available. Install with: pip install scikit-learn")
    SKLEARN_AVAILABLE = False

# 导入新的代码分析器和智能分块器
from backend.core.code_analyzer import CodeAnalyzer, analyze_code
from backend.core.smart_chunker import SmartChunker, chunk_content
from backend.core.document_summarizer import DocumentSummarizer, summarize_document

# 强制禁用 ChromaDB，只使用 TF-IDF
CHROMADB_AVAILABLE = False  # 强制禁用 ChromaDB


# 模型缓存
_model_cache = {}


def read_file_content(file_path: str, extract_images: bool = False) -> Dict[str, Any]:
    """
    读取文件内容，支持多种格式（包括 Word 文档）
    
    Args:
        file_path: 文件路径
        extract_images: 是否提取图片
        
    Returns:
        包含内容和图片的字典
    """
    ext = os.path.splitext(file_path)[1].lower()
    images = []
    
    try:
        if ext == '.docx' and DOCX_AVAILABLE:
            # 读取 Word 文档
            doc = DocxDocument(file_path)
            content = []
            
            # 提取段落内容
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(' | '.join(row_text))
            
            # 提取图片
            if extract_images:
                import io
                import base64
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_data = rel.target_part.blob
                            # 转换为 base64
                            image_base64 = base64.b64encode(image_data).decode('utf-8')
                            # 尝试获取图片类型
                            image_type = rel.target_ref.split('.')[-1].lower()
                            if image_type not in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
                                image_type = 'png'
                            
                            images.append({
                                "data": image_base64,
                                "type": image_type,
                                "description": f"[图片: {rel.target_ref}]"
                            })
                            
                            # 在内容中添加图片占位符
                            content.append(f"[图片: {rel.target_ref}]")
                        except Exception as e:
                            logger.warning(f"Failed to extract image: {e}")
            
            return {
                "content": '\n\n'.join(content),
                "images": images,
                "has_images": len(images) > 0
            }
        else:
            # 读取普通文本文件
            with open(file_path, 'r', encoding='utf-8', errors='strict') as f:
                text_content = f.read()
            
            return {
                "content": text_content,
                "images": [],
                "has_images": False
            }
    except UnicodeDecodeError:
        # 如果严格解码失败，尝试忽略错误
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text_content = f.read()
        
        return {
            "content": text_content,
            "images": [],
            "has_images": False
        }
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {e}")
        raise


def get_embedding_model(model_name: str = "paraphrase-multilingual-MiniLM-L12-v2"):
    """
    获取或缓存的嵌入模型
    
    Args:
        model_name: 模型名称
        
    Returns:
        嵌入模型实例或 None
    """
    global _model_cache
    
    if not SENTENCE_TRANSFORMERS_AVAILABLE:
        return None
    
    if model_name in _model_cache:
        return _model_cache[model_name]
    
    try:
        model = SentenceTransformer(model_name)
        _model_cache[model_name] = model
        logger.info(f"Loaded and cached embedding model: {model_name}")
        return model
    except Exception as e:
        logger.error(f"Failed to load embedding model {model_name}: {e}")
        return None


class Document:
    """文档类"""
    def __init__(
        self,
        content: str,
        metadata: Dict[str, Any],
        doc_id: str = None
    ):
        self.content = content
        self.metadata = metadata
        self.doc_id = doc_id or self._generate_id()
    
    def _generate_id(self) -> str:
        """生成文档 ID"""
        content_hash = hashlib.md5(self.content.encode()).hexdigest()[:16]
        return f"doc_{content_hash}"
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            "id": self.doc_id,
            "content": self.content,
            "metadata": self.metadata
        }


class RAGIndexer:
    """RAG 索引器 - 负责文档索引和嵌入生成"""
    
    def __init__(
        self,
        project_path: str,
        embedding_model: str = "paraphrase-multilingual-MiniLM-L12-v2",
        chunk_size: int = 1000,
        chunk_overlap: int = 100,
        incremental: bool = True
    ):
        """
        初始化索引器
        
        Args:
            project_path: 项目路径
            embedding_model: 嵌入模型名称
            chunk_size: 文档分块大小
            chunk_overlap: 分块重叠大小
            incremental: 是否启用增量索引
        """
        self.project_path = project_path
        self.embedding_model_name = embedding_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.incremental = incremental
        
        self.embedding_model = None
        self._init_embedding_model()
        
        # 初始化代码分析器、智能分块器和文档摘要器
        self.code_analyzer = CodeAnalyzer()
        self.smart_chunker = SmartChunker(chunk_size, 200, chunk_overlap)
        self.document_summarizer = DocumentSummarizer()
        
        # 支持的文件类型
        self.supported_extensions = {
            '.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.go', '.rs',
            '.md', '.txt', '.rst', '.json', '.yaml', '.yml',
            '.html', '.css', '.scss', '.less',
            '.docx'  # Word 文档支持
        }
        
        # 忽略的目录
        self.ignore_dirs = {
            'node_modules', '__pycache__', '.git', '.vscode',
            'dist', 'build', 'target', 'venv', 'env', '.env'
        }
        
        # 忽略的文件
        self.ignore_files = {
            '.gitignore', '.env', '.DS_Store', 'package-lock.json',
            'yarn.lock', 'pnpm-lock.yaml'
        }
        
        # 文件哈希缓存（用于增量索引）
        self.file_hashes = {}
        self._load_file_hashes()
    
    def _init_embedding_model(self):
        """初始化嵌入模型"""
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            logger.warning("Sentence transformers not available, embedding generation will be disabled")
            return
        
        try:
            # 使用多语言模型支持中英文
            self.embedding_model = SentenceTransformer(self.embedding_model_name)
            logger.info(f"Embedding model loaded: {self.embedding_model_name}")
        except Exception as e:
            logger.error(f"Failed to load embedding model: {e}")
    
    def _load_file_hashes(self):
        """加载文件哈希缓存"""
        hash_file = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            "..", "storage", "rag",
            hashlib.md5(self.project_path.encode()).hexdigest(),
            "file_hashes.json"
        )
        
        if os.path.exists(hash_file):
            try:
                with open(hash_file, 'r', encoding='utf-8') as f:
                    self.file_hashes = json.load(f)
                logger.info(f"Loaded {len(self.file_hashes)} file hashes")
            except Exception as e:
                logger.error(f"Failed to load file hashes: {e}")
    
    def _save_file_hashes(self):
        """保存文件哈希缓存"""
        hash_file = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            "..", "storage", "rag",
            hashlib.md5(self.project_path.encode()).hexdigest(),
            "file_hashes.json"
        )
        
        os.makedirs(os.path.dirname(hash_file), exist_ok=True)
        
        try:
            with open(hash_file, 'w', encoding='utf-8') as f:
                json.dump(self.file_hashes, f, indent=2)
            logger.info(f"Saved {len(self.file_hashes)} file hashes")
        except Exception as e:
            logger.error(f"Failed to save file hashes: {e}")
    
    def _get_file_hash(self, file_path: str) -> str:
        """计算文件哈希"""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception as e:
            logger.error(f"Failed to calculate hash for {file_path}: {e}")
            return ""
    
    def _is_file_changed(self, file_path: str) -> bool:
        """检查文件是否已更改"""
        if not self.incremental:
            return True
        
        rel_path = os.path.relpath(file_path, self.project_path)
        current_hash = self._get_file_hash(file_path)
        
        if not current_hash:
            return True
        
        if rel_path not in self.file_hashes:
            return True
        
        return self.file_hashes[rel_path] != current_hash
    
    def _split_text(self, text: str) -> List[str]:
        """将文本分割成块"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            # 计算结束位置
            end = start + self.chunk_size
            
            # 如果不是最后一块，尝试在句号或换行符处分割
            if end < len(text):
                # 寻找最近的句号
                last_period = text.rfind('。', start, end)
                last_dot = text.rfind('.', start, end)
                last_newline = text.rfind('\n', start, end)
                
                split_pos = max(last_period, last_dot, last_newline)
                if split_pos > start:
                    end = split_pos + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # 移动到下一块（带重叠）
            start = end - self.chunk_overlap
        
        return chunks
    
    def _extract_code_structure(self, content: str, file_path: str) -> Dict[str, Any]:
        """提取代码结构信息"""
        structure = {
            "functions": [],
            "classes": [],
            "imports": []
        }
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.py':
            self._extract_python_structure(content, structure)
        elif ext in ['.js', '.jsx', '.ts', '.tsx']:
            self._extract_javascript_structure(content, structure)
        
        return structure
    
    def _extract_python_structure(self, content: str, structure: Dict[str, Any]):
        """提取 Python 代码结构"""
        lines = content.split('\n')
        
        for line in lines:
            # 提取函数定义
            if line.strip().startswith('def '):
                func_name = line.strip().split('(')[0].replace('def ', '').strip()
                structure['functions'].append(func_name)
            
            # 提取类定义
            elif line.strip().startswith('class '):
                class_name = line.strip().split('(')[0].split(':')[0].replace('class ', '').strip()
                structure['classes'].append(class_name)
            
            # 提取导入
            elif line.strip().startswith('import ') or line.strip().startswith('from '):
                structure['imports'].append(line.strip())
    
    def _extract_javascript_structure(self, content: str, structure: Dict[str, Any]):
        """提取 JavaScript/TypeScript 代码结构"""
        lines = content.split('\n')
        
        for line in lines:
            # 提取函数定义
            if 'function ' in line or '=> ' in line:
                structure['functions'].append(line.strip()[:80])
            
            # 提取类定义
            elif 'class ' in line and line.strip().startswith('class '):
                class_name = line.strip().split('{')[0].replace('class ', '').strip()
                structure['classes'].append(class_name)
            
            # 提取导入
            elif line.strip().startswith('import ') or line.strip().startswith('require('):
                structure['imports'].append(line.strip())
    
    def _should_ignore_file(self, file_path: str) -> bool:
        """判断是否应该忽略文件"""
        filename = os.path.basename(file_path)
        
        # 检查文件扩展名
        ext = os.path.splitext(filename)[1].lower()
        if ext not in self.supported_extensions:
            return True
        
        # 检查文件名
        if filename in self.ignore_files:
            return True
        
        # 检查文件大小（限制 1MB）
        try:
            if os.path.getsize(file_path) > 1024 * 1024:
                logger.warning(f"File too large, skipping: {file_path}")
                return True
        except:
            return True
        
        # 检查是否为二进制文件
        try:
            with open(file_path, 'rb') as f:
                chunk = f.read(1024)
                # 检查是否包含大量非文本字符
                non_text_ratio = sum(1 for byte in chunk if byte > 127) / len(chunk) if chunk else 0
                if non_text_ratio > 0.3:  # 如果超过30%的字符是非文本字符，认为是二进制文件
                    logger.warning(f"Binary file detected, skipping: {file_path}")
                    return True
        except:
            return True
        
        return False
    
    async def index_project(self, progress_callback=None) -> List[Document]:
        """
        索引整个项目
        
        Args:
            progress_callback: 进度回调函数
        
        Returns:
            索引的文档列表
        """
        documents = []
        total_files = 0
        processed_files = 0
        
        # 统计文件数量
        for root, dirs, files in os.walk(self.project_path):
            # 过滤忽略的目录
            dirs[:] = [d for d in dirs if d not in self.ignore_dirs]
            
            for file in files:
                file_path = os.path.join(root, file)
                if not self._should_ignore_file(file_path):
                    total_files += 1
        
        logger.info(f"Found {total_files} files to index")
        
        # 遍历项目文件
        for root, dirs, files in os.walk(self.project_path):
            # 过滤忽略的目录
            dirs[:] = [d for d in dirs if d not in self.ignore_dirs]
            
            for file in files:
                file_path = os.path.join(root, file)
                
                if self._should_ignore_file(file_path):
                    continue
                
                try:
                    # 读取文件内容
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    
                    if not content.strip():
                        continue
                    
                    # 提取相对路径
                    rel_path = os.path.relpath(file_path, self.project_path)
                    
                    # 提取代码结构
                    structure = self._extract_code_structure(content, file_path)
                    
                    # 生成文档摘要
                    summary = self.document_summarizer.summarize(file_path, content)
                    
                    # 创建元数据（ChromaDB 只接受基本类型，字典需要转换为 JSON 字符串）
                    metadata = {
                        "file_path": rel_path,
                        "file_type": os.path.splitext(file)[1].lower(),
                        "file_size": len(content),
                        "indexed_at": datetime.now().isoformat(),
                        "structure": json.dumps(structure, ensure_ascii=False),
                        "summary": summary.get("summary", ""),
                        "language": summary.get("language", ""),
                        "total_lines": summary.get("total_lines", 0)
                    }
                    
                    # 分割文本
                    chunks = self._split_text(content)
                    
                    # 为每个块创建文档
                    for i, chunk in enumerate(chunks):
                        chunk_metadata = metadata.copy()
                        chunk_metadata["chunk_index"] = i
                        chunk_metadata["total_chunks"] = len(chunks)
                        
                        doc = Document(
                            content=chunk,
                            metadata=chunk_metadata
                        )
                        documents.append(doc)
                    
                    processed_files += 1
                    
                    # 调用进度回调
                    if progress_callback:
                        await progress_callback(processed_files, total_files, rel_path)
                    
                except Exception as e:
                    logger.error(f"Error indexing file {file_path}: {e}")
        
        logger.info(f"Indexed {len(documents)} chunks from {processed_files} files")
        return documents
    
    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """生成文本嵌入"""
        if not self.embedding_model:
            logger.warning("Embedding model not available")
            return []
        
        try:
            embeddings = self.embedding_model.encode(
                texts,
                show_progress_bar=False,
                convert_to_numpy=True
            )
            return embeddings.tolist()
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            return []


class TFIDFRetriever:
    """
    基于 TF-IDF 的轻量级检索器
    当 ChromaDB 不可用时的备选方案
    """
    
    def __init__(self, project_path: str):
        """
        初始化 TF-IDF 检索器
        
        Args:
            project_path: 项目路径
        """
        self.project_path = project_path
        self.vectorizer = None
        self.documents = []
        self.embeddings = None
        
        # 存储目录
        self.storage_dir = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            "..", "storage", "rag",
            hashlib.md5(project_path.encode()).hexdigest()
        )
        os.makedirs(self.storage_dir, exist_ok=True)
        
        # 加载已保存的索引
        self._load_index()
    
    def _load_index(self):
        """加载已保存的索引"""
        import pickle
        
        index_file = os.path.join(self.storage_dir, "tfidf_index.pkl")
        
        if os.path.exists(index_file):
            try:
                with open(index_file, 'rb') as f:
                    data = pickle.load(f)
                    self.vectorizer = data.get('vectorizer')
                    self.documents = data.get('documents', [])
                    self.embeddings = data.get('embeddings')
                logger.info(f"Loaded TF-IDF index with {len(self.documents)} documents")
            except Exception as e:
                logger.error(f"Failed to load TF-IDF index: {e}")
    
    def _save_index(self):
        """保存索引到磁盘"""
        import pickle
        
        index_file = os.path.join(self.storage_dir, "tfidf_index.pkl")
        
        try:
            with open(index_file, 'wb') as f:
                pickle.dump({
                    'vectorizer': self.vectorizer,
                    'documents': self.documents,
                    'embeddings': self.embeddings
                }, f)
            logger.info(f"Saved TF-IDF index with {len(self.documents)} documents")
        except Exception as e:
            logger.error(f"Failed to save TF-IDF index: {e}")
    
    async def add_documents(
        self,
        documents: List[Document],
        embeddings: List[List[float]] = None,
        progress_callback=None
    ):
        """
        添加文档到索引
        
        Args:
            documents: 文档列表
            embeddings: 预计算的嵌入（忽略，使用 TF-IDF）
            progress_callback: 进度回调函数
        """
        if not documents:
            return
        
        # 添加新文档
        self.documents.extend(documents)
        
        # 重新构建 TF-IDF 向量
        texts = [doc.content for doc in self.documents]
        self.vectorizer = TfidfVectorizer(
            max_features=5000,
            stop_words='english',
            ngram_range=(1, 2)
        )
        self.embeddings = self.vectorizer.fit_transform(texts)
        
        # 保存索引
        self._save_index()
        
        logger.info(f"Added {len(documents)} documents to TF-IDF index, total: {len(self.documents)}")
    
    def retrieve(
        self,
        query: str,
        n_results: int = 5,
        filters: Dict[str, Any] = None
    ) -> List[Dict[str, Any]]:
        """
        检索相关文档
        
        Args:
            query: 查询文本
            n_results: 返回结果数量
            filters: 元数据过滤条件
        
        Returns:
            检索结果列表
        """
        if not self.vectorizer or not self.documents:
            return []
        
        try:
            # 转换查询为向量
            query_vector = self.vectorizer.transform([query])
            
            # 计算相似度
            similarities = cosine_similarity(query_vector, self.embeddings).flatten()
            
            # 获取最相似的文档
            top_indices = similarities.argsort()[-n_results:][::-1]
            
            # 构建结果
            results = []
            logger.info(f"TF-IDF 检索开始: 查询='{query}', 请求结果数={n_results}, 总文档数={len(self.documents)}")
            logger.info(f"  相似度数组: {similarities[:10]}... (前10个)")
            
            for idx in top_indices:
                doc = self.documents[idx]
                similarity = similarities[idx]
                
                # 应用过滤条件
                if filters:
                    match = True
                    for key, value in filters.items():
                        if doc.metadata.get(key) != value:
                            match = False
                            break
                    if not match:
                        continue
                
                # 添加详细调试日志
                logger.info(f"=" * 80)
                logger.info(f"TF-IDF 检索结果 #{len(results)+1}:")
                logger.info(f"  索引位置: {idx}")
                logger.info(f"  文档ID: {doc.doc_id}")
                logger.info(f"  文件路径: {doc.metadata.get('file_path')}")
                logger.info(f"  块索引: {doc.metadata.get('chunk_index')}/{doc.metadata.get('total_chunks')}")
                logger.info(f"  相似度: {similarity:.6f}")
                logger.info(f"  距离: {1 - similarity:.6f}")
                logger.info(f"  行号范围: {doc.metadata.get('start_line')}-{doc.metadata.get('end_line')}")
                logger.info(f"  语言: {doc.metadata.get('language', 'N/A')}")
                logger.info(f"  内容长度: {len(doc.content)} 字符")
                logger.info(f"  内容预览: {doc.content[:200]}")
                logger.info(f"  完整元数据: {doc.metadata}")
                logger.info(f"=" * 80)
                
                results.append({
                    "id": doc.doc_id,
                    "content": doc.content,
                    "metadata": doc.metadata,
                    "distance": 1 - similarity,  # 转换为距离
                    "similarity": similarity
                })
            
            logger.info(f"TF-IDF 检索完成: 返回 {len(results)} 个结果")
            
            return results
        
        except Exception as e:
            logger.error(f"Error retrieving documents: {e}")
            return []
    
    def delete_collection(self):
        """删除集合"""
        import shutil
        
        try:
            if os.path.exists(self.storage_dir):
                shutil.rmtree(self.storage_dir)
            self.vectorizer = None
            self.documents = []
            self.embeddings = None
            logger.info(f"Deleted TF-IDF index")
        except Exception as e:
            logger.error(f"Error deleting TF-IDF index: {e}")
    
    def get_stats(self) -> Dict[str, Any]:
        """获取集合统计信息"""
        try:
            return {
                "collection_name": "tfidf_index",
                "document_count": len(self.documents),
                "project_path": self.project_path,
                "retriever_type": "tfidf"
            }
        except Exception as e:
            logger.error(f"Error getting stats: {e}")
            return {"error": str(e)}


class RAGRetriever:
    """RAG 检索器 - 负责检索相关文档"""
    
    def __init__(
        self,
        project_path: str,
        chroma_persist_dir: str = None
    ):
        """
        初始化检索器
        
        Args:
            project_path: 项目路径
            chroma_persist_dir: ChromaDB 持久化目录
        """
        self.project_path = project_path
        
        if not CHROMADB_AVAILABLE:
            raise RuntimeError("ChromaDB not available")
        
        # 设置持久化目录
        if chroma_persist_dir is None:
            chroma_persist_dir = os.path.join(
                os.path.dirname(os.path.dirname(__file__)),
                "..", "storage", "rag",
                hashlib.md5(project_path.encode()).hexdigest()
            )
        
        os.makedirs(chroma_persist_dir, exist_ok=True)
        
        # 初始化 ChromaDB
        self.client = chromadb.PersistentClient(
            path=chroma_persist_dir,
            settings=Settings(
                anonymized_telemetry=False,
                allow_reset=True
            )
        )
        
        # 获取或创建集合
        self.collection_name = "documents"
        try:
            self.collection = self.client.get_collection(name=self.collection_name)
            logger.info(f"Loaded existing collection: {self.collection_name}")
        except:
            self.collection = self.client.create_collection(
                name=self.collection_name,
                metadata={"project_path": project_path}
            )
            logger.info(f"Created new collection: {self.collection_name}")
    
    async def add_documents(
        self,
        documents: List[Document],
        embeddings: List[List[float]] = None,
        progress_callback=None
    ):
        """
        添加文档到向量数据库
        
        Args:
            documents: 文档列表
            embeddings: 预计算的嵌入（可选）
            progress_callback: 进度回调函数
        """
        if not documents:
            return
        
        # 批量添加
        batch_size = 100
        total = len(documents)
        
        for i in range(0, total, batch_size):
            batch = documents[i:i + batch_size]
            
            ids = [doc.doc_id for doc in batch]
            texts = [doc.content for doc in batch]
            metadatas = [doc.metadata for doc in batch]
            
            # 如果没有提供嵌入，使用 ChromaDB 的默认嵌入
            if embeddings is None:
                self.collection.add(
                    ids=ids,
                    documents=texts,
                    metadatas=metadatas
                )
            else:
                batch_embeddings = embeddings[i:i + batch_size]
                self.collection.add(
                    ids=ids,
                    documents=texts,
                    metadatas=metadatas,
                    embeddings=batch_embeddings
                )
            
            if progress_callback:
                await progress_callback(min(i + batch_size, total), total)
        
        logger.info(f"Added {total} documents to collection")
    
    def retrieve(
        self,
        query: str,
        n_results: int = 5,
        filters: Dict[str, Any] = None
    ) -> List[Dict[str, Any]]:
        """
        检索相关文档
        
        Args:
            query: 查询文本
            n_results: 返回结果数量
            filters: 元数据过滤条件
        
        Returns:
            检索结果列表
        """
        try:
            # 构建查询参数
            query_params = {
                "query_texts": [query],
                "n_results": n_results
            }
            
            # 添加过滤条件
            if filters:
                query_params["where"] = filters
            
            # 执行查询
            results = self.collection.query(**query_params)
            
            # 格式化结果
            formatted_results = []
            if results['ids'] and results['ids'][0]:
                for i in range(len(results['ids'][0])):
                    formatted_results.append({
                        "id": results['ids'][0][i],
                        "content": results['documents'][0][i],
                        "metadata": results['metadatas'][0][i],
                        "distance": results['distances'][0][i] if 'distances' in results else None
                    })
            
            return formatted_results
        
        except Exception as e:
            logger.error(f"Error retrieving documents: {e}")
            return []
    
    def delete_collection(self):
        """删除集合"""
        try:
            self.client.delete_collection(name=self.collection_name)
            logger.info(f"Deleted collection: {self.collection_name}")
        except Exception as e:
            logger.error(f"Error deleting collection: {e}")
    
    def get_stats(self) -> Dict[str, Any]:
        """获取集合统计信息"""
        try:
            count = self.collection.count()
            return {
                "collection_name": self.collection_name,
                "document_count": count,
                "project_path": self.project_path
            }
        except Exception as e:
            logger.error(f"Error getting stats: {e}")
            return {"error": str(e)}


class HybridRetriever:
    """
    混合检索器 - 结合 BM25 和语义检索
    提供更准确的检索结果
    """
    
    def __init__(self, project_path: str):
        """
        初始化混合检索器
        
        Args:
            project_path: 项目路径
        """
        self.project_path = project_path
        
        # 使用 TF-IDF 作为 BM25 的近似
        if not SKLEARN_AVAILABLE:
            raise RuntimeError("scikit-learn not available for hybrid retrieval")
        
        self.tfidf_retriever = TFIDFRetriever(project_path)
        
        # 如果可用，使用 ChromaDB 进行语义检索
        self.semantic_retriever = None
        if CHROMADB_AVAILABLE:
            self.semantic_retriever = RAGRetriever(project_path)
    
    async def add_documents(
        self,
        documents: List[Document],
        embeddings: List[List[float]] = None,
        progress_callback=None
    ):
        """添加文档到两个检索器"""
        # 添加到 TF-IDF 检索器
        await self.tfidf_retriever.add_documents(documents, embeddings, progress_callback)
        
        # 添加到语义检索器（如果可用）
        if self.semantic_retriever:
            await self.semantic_retriever.add_documents(documents, embeddings, progress_callback)
    
    def retrieve(
        self,
        query: str,
        n_results: int = 5,
        filters: Dict[str, Any] = None,
        alpha: float = 0.5
    ) -> List[Dict[str, Any]]:
        """
        混合检索 - 结合关键词和语义检索
        
        Args:
            query: 查询文本
            n_results: 返回结果数量
            filters: 元数据过滤条件
            alpha: 语义检索权重 (0-1)，0.5 表示平衡
        
        Returns:
            检索结果列表
        """
        # 获取 TF-IDF 结果
        tfidf_results = self.tfidf_retriever.retrieve(query, n_results * 2, filters)
        
        # 获取语义检索结果（如果可用）
        semantic_results = []
        if self.semantic_retriever:
            semantic_results = self.semantic_retriever.retrieve(query, n_results * 2, filters)
        
        # 融合结果
        fused_results = self._fuse_results(
            tfidf_results,
            semantic_results,
            alpha,
            n_results
        )
        
        return fused_results
    
    def _fuse_results(
        self,
        tfidf_results: List[Dict],
        semantic_results: List[Dict],
        alpha: float,
        n_results: int
    ) -> List[Dict[str, Any]]:
        """
        融合两个检索结果
        
        使用倒数排名融合（Reciprocal Rank Fusion, RRF）
        """
        # 创建文档 ID 到分数的映射
        scores = {}
        
        # TF-IDF 分数（转换为排名分数）
        for rank, result in enumerate(tfidf_results):
            doc_id = result["id"]
            # RRF: 1 / (k + rank)，k 通常为 60
            rrf_score = 1.0 / (60 + rank + 1)
            scores[doc_id] = scores.get(doc_id, 0) + rrf_score * (1 - alpha)
        
        # 语义检索分数
        for rank, result in enumerate(semantic_results):
            doc_id = result["id"]
            rrf_score = 1.0 / (60 + rank + 1)
            scores[doc_id] = scores.get(doc_id, 0) + rrf_score * alpha
        
        # 按分数排序
        sorted_results = sorted(scores.items(), key=lambda x: x[1], reverse=True)
        
        # 获取前 n_results 个结果
        final_results = []
        result_map = {r["id"]: r for r in tfidf_results + semantic_results}
        
        for doc_id, score in sorted_results[:n_results]:
            if doc_id in result_map:
                result = result_map[doc_id].copy()
                result["fusion_score"] = score
                final_results.append(result)
        
        return final_results
    
    def delete_collection(self):
        """删除集合"""
        self.tfidf_retriever.delete_collection()
        if self.semantic_retriever:
            self.semantic_retriever.delete_collection()
    
    def get_stats(self) -> Dict[str, Any]:
        """获取统计信息"""
        stats = self.tfidf_retriever.get_stats()
        stats["retriever_type"] = "hybrid"
        if self.semantic_retriever:
            semantic_stats = self.semantic_retriever.get_stats()
            stats["semantic_available"] = True
        else:
            stats["semantic_available"] = False
        return stats


class RAGService:
    """RAG 服务 - 统一的 RAG 功能接口"""
    
    def __init__(self, project_path: str, use_chromadb: bool = None, use_hybrid: bool = False):
        """
        初始化 RAG 服务
        
        Args:
            project_path: 项目路径
            use_chromadb: 是否使用 ChromaDB（None 表示自动检测）
            use_hybrid: 是否使用混合检索
        """
        self.project_path = project_path
        self.indexer = RAGIndexer(project_path)
        self.retriever = None
        self._initialized = False
        self.use_hybrid = use_hybrid
        
        # 自动选择检索器
        if use_hybrid:
            # 使用混合检索
            self.use_chromadb = CHROMADB_AVAILABLE
            logger.info(f"RAG Service initialized with Hybrid retriever")
        elif use_chromadb is None:
            # 优先使用 ChromaDB，如果不可用则使用 TF-IDF
            self.use_chromadb = CHROMADB_AVAILABLE
            logger.info(f"RAG Service initialized with {'ChromaDB' if self.use_chromadb else 'TF-IDF'} retriever")
        else:
            self.use_chromadb = use_chromadb
            logger.info(f"RAG Service initialized with {'ChromaDB' if self.use_chromadb else 'TF-IDF'} retriever")
    
    def _ensure_initialized(self):
        """确保检索器已初始化"""
        if not self._initialized:
            logger.info(f"Initializing retriever: use_chromadb={self.use_chromadb}, use_hybrid={self.use_hybrid}")
            logger.info(f"CHROMADB_AVAILABLE={CHROMADB_AVAILABLE}, SKLEARN_AVAILABLE={SKLEARN_AVAILABLE}")
            
            if self.use_hybrid and SKLEARN_AVAILABLE:
                # 使用混合检索
                self.retriever = HybridRetriever(self.project_path)
                logger.info("Using HybridRetriever")
            elif self.use_chromadb and CHROMADB_AVAILABLE:
                self.retriever = RAGRetriever(self.project_path)
                logger.info("Using RAGRetriever (ChromaDB)")
            elif SKLEARN_AVAILABLE:
                self.retriever = TFIDFRetriever(self.project_path)
                logger.info("Using TFIDFRetriever (scikit-learn)")
            else:
                logger.error("No retriever available! Neither ChromaDB nor scikit-learn is installed.")
                return
            self._initialized = True
            logger.info(f"Retriever initialized successfully")
    
    async def index_project(self, progress_callback=None, force_reindex: bool = False) -> AsyncGenerator[Dict[str, Any], None]:
        """
        索引项目（支持增量索引）
        
        Args:
            progress_callback: 进度回调函数（已废弃，使用生成器返回进度）
            force_reindex: 是否强制重新索引所有文件
        
        Yields:
            进度更新字典
        """
        logger.info(f"Starting project indexing: {self.project_path} (force_reindex={force_reindex})")
        
        yield {"type": "status", "message": "开始扫描项目文件...", "progress": 0}
        
        # 获取文件列表
        all_files = []
        changed_files = []
        deleted_files = []
        
        for root, dirs, files in os.walk(self.project_path):
            # 过滤忽略的目录
            dirs[:] = [d for d in dirs if d in self.indexer.ignore_dirs]
            
            for file in files:
                file_path = os.path.join(root, file)
                
                if self.indexer._should_ignore_file(file_path):
                    continue
                
                all_files.append(file_path)
        
        yield {"type": "status", "message": f"发现 {len(all_files)} 个文件，检查变更...", "progress": 10}
        
        # 检查文件变更
        for file_path in all_files:
            rel_path = os.path.relpath(file_path, self.project_path)
            current_hash = self.indexer._get_file_hash(file_path)
            
            if force_reindex:
                # 强制重新索引所有文件
                changed_files.append((file_path, rel_path, current_hash))
            else:
                # 检查是否是新文件或已更改的文件
                old_hash = self.indexer.file_hashes.get(rel_path, "")
                if old_hash != current_hash:
                    changed_files.append((file_path, rel_path, current_hash))
        
        # 检查删除的文件
        for rel_path in list(self.indexer.file_hashes.keys()):
            full_path = os.path.join(self.project_path, rel_path)
            if not os.path.exists(full_path):
                deleted_files.append(rel_path)
        
        yield {
            "type": "status",
            "message": f"发现 {len(changed_files)} 个变更文件，{len(deleted_files)} 个删除文件",
            "progress": 20
        }
        
        if not changed_files and not deleted_files:
            yield {
                "type": "complete",
                "message": "没有文件变更，无需重新索引",
                "progress": 100,
                "stats": self.get_stats() if self._initialized else {}
            }
            return
        
        # 处理变更的文件
        documents = []
        processed_count = 0
        
        for file_path, rel_path, file_hash in changed_files:
            try:
                # 读取文件内容（使用新的 read_file_content 函数支持 Word 文档和图片）
                try:
                    file_data = read_file_content(file_path, extract_images=True)
                    content = file_data["content"]
                    images = file_data.get("images", [])
                except Exception as e:
                    logger.warning(f"Cannot read file, skipping: {file_path} - {e}")
                    continue
                
                if not content.strip():
                    continue
                
                # 对于非 Word 文档，检查内容是否包含过多不可打印字符（可能是二进制文件）
                ext = os.path.splitext(file_path)[1].lower()
                if ext != '.docx':
                    non_printable_ratio = sum(1 for c in content if ord(c) > 127) / len(content) if content else 0
                    if non_printable_ratio > 0.3:
                        logger.warning(f"File contains too many non-printable characters, skipping: {file_path}")
                        continue
                
                # 提取代码结构
                structure = self.indexer._extract_code_structure(content, file_path)
                
                # 生成文档摘要
                summary = self.indexer.document_summarizer.summarize(file_path, content)
                
                # 创建元数据（ChromaDB 只接受基本类型，字典需要转换为 JSON 字符串）
                metadata = {
                    "file_path": rel_path,
                    "file_type": os.path.splitext(file_path)[1].lower(),
                    "file_size": len(content),
                    "indexed_at": datetime.now().isoformat(),
                    "file_hash": file_hash,
                    "structure": json.dumps(structure, ensure_ascii=False),
                    "summary": summary.get("summary", ""),
                    "language": summary.get("language", ""),
                    "total_lines": summary.get("total_lines", 0)
                }
                
                # 使用智能分块器
                chunks = self.indexer.smart_chunker.chunk(content, file_path, structure)
                
                # 为每个块创建文档
                for i, chunk_data in enumerate(chunks):
                    chunk_metadata = metadata.copy()
                    chunk_metadata.update(chunk_data["metadata"])
                    chunk_metadata["chunk_index"] = i
                    chunk_metadata["total_chunks"] = len(chunks)
                    
                    doc = Document(
                        content=chunk_data["content"],
                        metadata=chunk_metadata
                    )
                    documents.append(doc)
                
                # 更新文件哈希
                self.indexer.file_hashes[rel_path] = file_hash
                processed_count += 1
                
                # 更新进度
                progress = 20 + int((processed_count / len(changed_files)) * 50)
                if processed_count % 10 == 0:  # 每10个文件更新一次
                    yield {
                        "type": "status",
                        "message": f"已处理 {processed_count}/{len(changed_files)} 个文件...",
                        "progress": progress
                    }
                
            except Exception as e:
                logger.error(f"Error indexing file {file_path}: {e}")
        
        # 保存文件哈希
        self.indexer._save_file_hashes()
        
        if not documents:
            yield {
                "type": "complete",
                "message": "没有新的文档需要索引",
                "progress": 100,
                "stats": self.get_stats() if self._initialized else {}
            }
            return
        
        yield {"type": "status", "message": f"生成 {len(documents)} 个文档块，正在添加到数据库...", "progress": 75}
        
        # 生成嵌入（仅在使用 ChromaDB 时）
        embeddings = None
        if self.use_chromadb and self.indexer.embedding_model:
            texts = [doc.content for doc in documents]
            embeddings = self.indexer.generate_embeddings(texts)
        
        # 添加到向量数据库
        self._ensure_initialized()
        await self.retriever.add_documents(documents, embeddings, progress_callback)
        
        # 获取统计信息
        stats = self.retriever.get_stats()
        
        yield {
            "type": "complete",
            "message": f"成功索引 {len(documents)} 个文档块 ({len(changed_files)} 个文件变更)",
            "progress": 100,
            "stats": stats,
            "changed_files": len(changed_files),
            "deleted_files": len(deleted_files)
        }
    
    def retrieve(self, query: str, n_results: int = 5) -> List[Dict[str, Any]]:
        """
        检索相关文档
        
        Args:
            query: 查询文本
            n_results: 返回结果数量
        
        Returns:
            检索结果
        """
        self._ensure_initialized()
        return self.retriever.retrieve(query, n_results)
    
    async def add_document(
        self,
        file_name: str,
        content: str,
        file_type: str = None,
        images: List[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        添加单个文档到索引
        
        Args:
            file_name: 文件名
            content: 文档内容
            file_type: 文件类型（可选）
            images: 图片列表（可选）
        
        Returns:
            添加结果
        """
        if not content.strip():
            return {"success": False, "error": "文档内容为空"}
        
        if images is None:
            images = []
        
        # 确定文件类型
        if file_type is None:
            file_type = os.path.splitext(file_name)[1].lower()
        
        # 提取代码结构
        structure = self.indexer._extract_code_structure(content, file_name)
        
        # 生成文档摘要
        summary = self.indexer.document_summarizer.summarize(file_name, content)
        
        # 创建元数据
        file_hash = hashlib.md5(content.encode()).hexdigest()
        metadata = {
            "file_path": f"uploaded/{file_name}",
            "file_type": file_type,
            "file_size": len(content),
            "indexed_at": datetime.now().isoformat(),
            "file_hash": file_hash,
            "structure": json.dumps(structure, ensure_ascii=False),
            "summary": summary.get("summary", ""),
            "language": summary.get("language", ""),
            "total_lines": summary.get("total_lines", 0),
            "uploaded": True,
            "has_images": len(images) > 0,
            "image_count": len(images)
        }
        
        # 使用智能分块器
        chunks = self.indexer.smart_chunker.chunk(content, file_name, structure)
        
        # 为每个块创建文档
        documents = []
        for i, chunk_data in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata.update(chunk_data["metadata"])
            chunk_metadata["chunk_index"] = i
            chunk_metadata["total_chunks"] = len(chunks)
            
            # 如果有图片，在内容中添加图片信息
            chunk_content = chunk_data["content"]
            if images:
                # 在文档开头添加图片信息
                image_info = "\n\n".join([img.get("description", "[图片]") for img in images])
                chunk_content = f"[文档包含 {len(images)} 张图片]\n{image_info}\n\n{chunk_content}"
            
            doc = Document(
                content=chunk_content,
                metadata=chunk_metadata
            )
            documents.append(doc)
        
        # 添加到向量数据库
        self._ensure_initialized()
        await self.retriever.add_documents(documents)
        
        return {
            "success": True,
            "message": f"成功添加文档: {file_name}",
            "chunks": len(documents),
            "file_name": file_name
        }
    
    async def add_documents_from_files(
        self,
        file_paths: List[str]
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """
        从文件路径批量添加文档
        
        Args:
            file_paths: 文件路径列表
        
        Yields:
            进度更新
        """
        total = len(file_paths)
        processed = 0
        
        for file_path in file_paths:
            try:
                # 读取文件（使用新的 read_file_content 函数支持 Word 文档和图片）
                try:
                    file_data = read_file_content(file_path, extract_images=True)
                    content = file_data["content"]
                    images = file_data.get("images", [])
                except Exception as e:
                    logger.warning(f"Cannot read file, skipping: {file_path} - {e}")
                    yield {
                        "type": "progress",
                        "file": file_path,
                        "status": "error",
                        "message": f"无法读取文件: {str(e)}"
                    }
                    continue
                
                file_name = os.path.basename(file_path)
                file_type = os.path.splitext(file_name)[1].lower()
                
                # 添加文档
                result = await self.add_document(file_name, content, file_type, images=images)
                
                processed += 1
                
                yield {
                    "type": "progress",
                    "file_name": file_name,
                    "processed": processed,
                    "total": total,
                    "success": result.get("success", False),
                    "message": result.get("message", "")
                }
                
            except Exception as e:
                logger.error(f"Error adding document {file_path}: {e}")
                processed += 1
                
                yield {
                    "type": "error",
                    "file_name": os.path.basename(file_path),
                    "processed": processed,
                    "total": total,
                    "error": str(e)
                }
        
        yield {
            "type": "complete",
            "message": f"成功处理 {processed} 个文件",
            "total": total
        }
    
    def get_stats(self) -> Dict[str, Any]:
        """获取统计信息"""
        self._ensure_initialized()
        return self.retriever.get_stats()
    
    def reset(self):
        """重置 RAG 服务"""
        if self.retriever:
            self.retriever.delete_collection()
        self._initialized = False


# 便捷函数
def get_rag_service(project_path: str, use_chromadb: bool = None) -> RAGService:
    """
    获取 RAG 服务实例
    
    Args:
        project_path: 项目路径
        use_chromadb: 是否使用 ChromaDB（None 表示自动检测）
    
    Returns:
        RAG 服务实例
    """
    return RAGService(project_path, use_chromadb)
